from docx import Document
import os

def create_report(ddr_text, inspection_images, thermal_images, output_file):

    doc = Document()

    doc.add_heading("Detailed Diagnostic Report (DDR)", level=1)

    doc.add_paragraph(ddr_text)

    doc.add_heading("Inspection Images", level=2)

    for img in os.listdir(inspection_images):

        img_path = os.path.join(inspection_images, img)
        doc.add_picture(img_path)

    doc.add_heading("Thermal Images", level=2)

    for img in os.listdir(thermal_images):

        img_path = os.path.join(thermal_images, img)
        doc.add_picture(img_path)

    doc.save(output_file)